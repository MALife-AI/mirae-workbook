#!/usr/bin/env python3
"""handout.md → handout.pdf + handout.docx.

PDF는 책처럼 양면(짝수=왼쪽, 홀수=오른쪽)으로 펼쳐서 볼 수 있도록
안쪽(여는 면) 여백을 더 주고 줄간격·바깥 여백을 넉넉히 잡았다.
"""
from pathlib import Path
import markdown
from weasyprint import HTML, CSS

DOCS = Path(__file__).parent
MD = DOCS / "handout.md"
PDF = DOCS / "handout.pdf"
DOCX = DOCS / "handout.docx"

html_body = markdown.markdown(
    MD.read_text(encoding="utf-8"),
    extensions=["tables", "fenced_code", "toc"],
)

# ──────────────────────────────────────────────────────────────────
# 목차 자동 생성 — 헤딩 트리에서 추출 → 페이지 번호 자동 삽입
# ──────────────────────────────────────────────────────────────────
def _build_auto_toc(html):
    """기존 수기 목차를 헤딩 기반 자동 목차로 교체.

    PDF: <a> 안에 빈 ::after 가 target-counter(attr(href), page) 로 채워짐.
    DOCX: 빌더가 .auto-toc 컨테이너를 인식해 Word TOC 필드로 치환.
    """
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # 모든 h1/h2 수집 (문서 제목과 목차 자체 제외)
    entries = []
    skip_titles = {"목차", "Claude Code 워크북"}
    for tag in soup.find_all(["h1", "h2"]):
        text = tag.get_text().strip()
        if text in skip_titles:
            continue
        entries.append((tag.name, tag.get("id", ""), text))

    # 자동 목차 HTML 만들기
    toc_lines = ['<div class="auto-toc">']
    for level, anchor, text in entries:
        cls = "toc-h1" if level == "h1" else "toc-h2"
        toc_lines.append(
            f'<p class="{cls}"><a href="#{anchor}">{text}</a></p>'
        )
    toc_lines.append("</div>")
    new_toc_html = "\n".join(toc_lines)

    # "목차" h1 ~ 다음 h1 (Part 1) 사이의 노드를 제거하고 새 목차로 교체
    toc_h1 = None
    for h1 in soup.find_all("h1"):
        if h1.get_text().strip() == "목차":
            toc_h1 = h1
            break
    if toc_h1 is None:
        return html

    # 다음 h1 까지의 sibling 들을 모두 제거
    next_node = toc_h1.next_sibling
    to_remove = []
    while next_node is not None:
        if getattr(next_node, "name", None) == "h1":
            break
        to_remove.append(next_node)
        next_node = next_node.next_sibling
    for n in to_remove:
        if hasattr(n, "decompose"):
            n.decompose()
        else:
            n.extract()

    # 새 목차 노드 삽입
    new_fragment = BeautifulSoup(new_toc_html, "html.parser")
    toc_h1.insert_after(new_fragment)
    return str(soup)


html_body = _build_auto_toc(html_body)

# ──────────────────────────────────────────────────────────────────
# PDF — 책처럼 양면 펼침 레이아웃
# ──────────────────────────────────────────────────────────────────
CSS_STR = """
/* 책 펼침: 왼쪽(짝수)·오른쪽(홀수) 페이지의 안쪽 여백을 넓게 */
@page {
  size: A4;
  margin: 22mm 18mm 22mm 18mm;
}
@page :left {
  margin-left: 18mm;   /* 바깥쪽 */
  margin-right: 26mm;  /* 안쪽(제본) */
  @bottom-left  { content: counter(page); font-size: 9pt; color: #888; }
  @top-left     { content: "미래에셋생명 · Vibe Coding 워크북"; font-size: 8pt; color: #bbb; }
}
@page :right {
  margin-left: 26mm;   /* 안쪽(제본) */
  margin-right: 18mm;  /* 바깥쪽 */
  @bottom-right { content: counter(page); font-size: 9pt; color: #888; }
  @top-right    { content: string(chapter); font-size: 8pt; color: #bbb; }
}
@page :first {
  margin: 28mm 22mm;
  @top-left { content: ""; }
  @top-right { content: ""; }
  @bottom-left { content: ""; }
  @bottom-right { content: ""; }
}

html, body {
  font-family: "Noto Sans KR", sans-serif;
  font-size: 10.5pt;
  line-height: 1.75;       /* 책처럼 넉넉한 줄간격 */
  color: #1a1a1a;
}

/* 각 챕터(h1)는 항상 오른쪽 페이지(홀수)에서 시작 — 책처럼 */
h1 {
  font-size: 22pt;
  color: #F58220;
  border-bottom: 2px solid #F58220;
  padding-bottom: 6pt;
  margin-top: 0;
  margin-bottom: 14pt;
  page-break-before: right;   /* 새 챕터는 오른쪽 페이지부터 */
  page-break-after: avoid;
  string-set: chapter content();
}
h1:first-of-type { page-break-before: avoid; }

h2 {
  font-size: 15pt;
  color: #1a1a1a;
  margin-top: 18pt;
  margin-bottom: 8pt;
  border-left: 4px solid #F58220;
  padding-left: 10px;
  page-break-after: avoid;
}
h3 {
  font-size: 12pt;
  color: #9a3412;
  margin-top: 14pt;
  margin-bottom: 5pt;
  page-break-after: avoid;
}
h4 {
  font-size: 11pt;
  color: #555;
  margin-top: 10pt;
  margin-bottom: 4pt;
  page-break-after: avoid;
}
p { margin: 7pt 0; }
strong { color: #1a1a1a; font-weight: 700; }
em { color: #9a3412; font-style: italic; }

code {
  font-family: "Noto Sans KR", monospace;
  background: #f5f5f5;
  padding: 1pt 4pt;
  border-radius: 3px;
  font-size: 9.5pt;
  color: #c2410c;
}
pre {
  background: #1e1e1e;
  color: #f5f5f5;
  padding: 10pt 12pt;
  border-radius: 5px;
  font-size: 9pt;
  overflow: hidden;
  line-height: 1.55;
  margin: 10pt 0;
  page-break-inside: avoid;
}
pre code { background: none; color: inherit; padding: 0; font-size: inherit; }

table {
  border-collapse: collapse;
  width: 100%;
  margin: 10pt 0;
  font-size: 9.5pt;
  page-break-inside: avoid;
}
th, td {
  border: 1px solid #ddd;
  padding: 6pt 8pt;
  text-align: left;
  vertical-align: top;
}
th { background: #fff7ed; color: #9a3412; font-weight: 700; }

ul, ol { margin: 8pt 0; padding-left: 22pt; }
li { margin: 4pt 0; line-height: 1.7; }

hr {
  border: none;
  border-top: 1px solid #F58220;
  margin: 16pt 0;
}
blockquote {
  border-left: 3px solid #F58220;
  padding: 6pt 14pt;
  margin: 10pt 0;
  background: #fff7ed;
  color: #555;
  font-size: 9.5pt;
  page-break-inside: avoid;
}
a { color: #F58220; text-decoration: none; }

/* 미아 줄·과부 줄 방지 */
p, li, blockquote { orphans: 3; widows: 3; }

/* ── 자동 목차 — leader dot + 페이지 번호 ── */
.auto-toc { margin-top: 14pt; }
.auto-toc p {
  margin: 0;
  padding: 5pt 0;
  border-bottom: 0.5pt dotted #e5e5e5;
  page-break-inside: avoid;       /* 한 항목이 두 페이지에 걸치지 않게 */
  break-inside: avoid;
}
.auto-toc .toc-h1 {
  margin-top: 14pt;
  font-size: 12pt;
  font-weight: 700;
  color: #F58220;
  border-bottom: 1pt solid #F58220;
  page-break-after: avoid;        /* Part 헤더 직후 첫 항목과 떨어지지 않게 */
  break-after: avoid;
}
.auto-toc .toc-h2 {
  margin-left: 14pt;
  font-size: 10.5pt;
  color: #1a1a1a;
}
.auto-toc a { color: inherit; text-decoration: none; display: block; }
.auto-toc a::after {
  content: leader('.') target-counter(attr(href), page);
  color: #888;
  font-weight: 400;
}
"""

html_doc = f"""<!DOCTYPE html>
<html lang="ko"><head><meta charset="utf-8"><title>Claude Code 워크북</title></head>
<body>{html_body}</body></html>
"""

HTML(string=html_doc).write_pdf(target=str(PDF), stylesheets=[CSS(string=CSS_STR)])
print(f"OK → {PDF} ({PDF.stat().st_size // 1024} KB)")

# ──────────────────────────────────────────────────────────────────
# DOCX — PDF 디자인을 그대로 구현 (python-docx + bs4)
# ──────────────────────────────────────────────────────────────────
BRAND_OR = "F58220"
BRAND_OR_DARK = "9A3412"
BRAND_OR_LIGHT = "FFF7ED"
TX_DARK = "1A1A1A"
TX_GRAY = "555555"
CODE_BG_DARK = "1E1E1E"
CODE_FG_LIGHT = "F5F5F5"
INLINE_CODE_BG = "F5F5F5"
INLINE_CODE_FG = "C2410C"
KO_FONT = "Noto Sans KR"
MONO_FONT = "Consolas"


def _qn(tag):
    from docx.oxml.ns import qn
    return qn(tag)


def _set_cell_shading(cell, hex_color):
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_paragraph_shading(paragraph, hex_color):
    from docx.oxml import OxmlElement
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_color)
    p_pr.append(shd)


def _set_paragraph_border(paragraph, *, left=None, bottom=None):
    """left/bottom: dict(size, color) or None. size = 1/8 pt units."""
    from docx.oxml import OxmlElement
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for name, opts in (("left", left), ("bottom", bottom)):
        if not opts:
            continue
        b = OxmlElement(f"w:{name}")
        b.set(_qn("w:val"), "single")
        b.set(_qn("w:sz"), str(opts.get("size", 24)))
        b.set(_qn("w:space"), str(opts.get("space", 6)))
        b.set(_qn("w:color"), opts.get("color", BRAND_OR))
        pbdr.append(b)
    p_pr.append(pbdr)


def _ensure_korean_font(run, font_name=KO_FONT):
    """동아시아 글꼴까지 확실히 적용 (한글 표시 보장)."""
    from docx.oxml import OxmlElement
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(_qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(_qn("w:ascii"), font_name)
    rFonts.set(_qn("w:hAnsi"), font_name)
    rFonts.set(_qn("w:eastAsia"), font_name)
    rFonts.set(_qn("w:cs"), font_name)


def _add_run(paragraph, text, *, bold=False, italic=False, color=TX_DARK,
             size=10.5, font=KO_FONT):
    from docx.shared import Pt, RGBColor
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    _ensure_korean_font(run, font)
    return run


def _walk_inline(paragraph, node, *, base_color=TX_DARK, base_size=10.5,
                 bold=False, italic=False):
    """HTML 인라인 노드를 paragraph 안의 run 들로 변환."""
    from bs4 import NavigableString
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            _add_run(paragraph, text, bold=bold, italic=italic,
                     color=base_color, size=base_size)
        return

    name = node.name
    if name in ("strong", "b"):
        for child in node.children:
            _walk_inline(paragraph, child, base_color=base_color,
                         base_size=base_size, bold=True, italic=italic)
    elif name in ("em", "i"):
        for child in node.children:
            _walk_inline(paragraph, child, base_color=BRAND_OR_DARK,
                         base_size=base_size, bold=bold, italic=True)
    elif name == "code":
        # 인라인 코드 — 배경/색상 강조
        text = node.get_text()
        run = _add_run(paragraph, text, bold=False, italic=False,
                       color=INLINE_CODE_FG, size=base_size - 0.5,
                       font=MONO_FONT)
        # 인라인 shading 은 run 단위가 까다롭다 — 색상만 강조하고 패스
    elif name == "a":
        text = node.get_text()
        _add_run(paragraph, text, bold=bold, italic=italic,
                 color=BRAND_OR, size=base_size)
    elif name == "br":
        paragraph.add_run().add_break()
    else:
        for child in node.children:
            _walk_inline(paragraph, child, base_color=base_color,
                         base_size=base_size, bold=bold, italic=italic)


def build_docx():
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── 페이지: A4 + 책 양면 (mirror) 여백
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)   # 안쪽(제본)
    section.right_margin = Cm(1.8)  # 바깥쪽
    # mirror margins 활성화
    sect_pr = section._sectPr
    from docx.oxml import OxmlElement
    page_mar = sect_pr.find(_qn("w:pgMar"))
    if page_mar is not None:
        page_mar.set(_qn("w:gutter"), "0")
    mirror = OxmlElement("w:mirrorMargins")
    sect_pr.insert(0, mirror)

    # ── 기본 스타일: Normal 폰트를 한글로
    normal = doc.styles["Normal"]
    normal.font.name = KO_FONT
    normal.font.size = Pt(10.5)
    normal_rpr = normal.element.get_or_add_rPr()
    rFonts = normal_rpr.find(_qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        normal_rpr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(_qn(attr), KO_FONT)
    normal.paragraph_format.line_spacing = 1.6
    normal.paragraph_format.space_after = Pt(4)

    soup = BeautifulSoup(html_body, "html.parser")

    first_h1 = True

    # 헤딩 → 추정 페이지 (DOCX는 실제 페이지 계산이 어려우므로
    # 본문 분량 비례로 추정 — Word TOC 필드로 갱신되기 전 폴백 표시)
    _toc_skip = {"목차", "Claude Code 워크북"}
    headings_seq = [t for t in soup.find_all(["h1", "h2"])
                    if t.get_text().strip() not in _toc_skip]

    def add_paragraph(style_name=None):
        p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        return p

    def add_word_toc_field():
        """Word TOC 필드 — 문서를 Word/한컴 등에서 열면 자동 갱신됨."""
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        run1 = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(_qn("w:fldCharType"), "begin")
        run1._r.append(fld_begin)

        run2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(_qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-2" \\h \\z \\u '
        run2._r.append(instr)

        run3 = p.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(_qn("w:fldCharType"), "separate")
        run3._r.append(fld_sep)

        run4 = p.add_run()
        _ensure_korean_font(run4)
        run4.text = "(Word·한컴오피스에서 'F9' 또는 우클릭 → 필드 업데이트로 페이지 번호가 갱신됩니다)"
        run4.font.size = Pt(9)
        from docx.shared import RGBColor
        run4.font.color.rgb = RGBColor.from_string("888888")
        run4.italic = True

        run5 = p.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(_qn("w:fldCharType"), "end")
        run5._r.append(fld_end)

    def add_static_toc_fallback():
        """Word TOC 필드 옆에 정적 TOC 도 함께 렌더 — 어디서든 보이게."""
        for h in headings_seq:
            level = h.name
            text = h.get_text().strip()
            p = doc.add_paragraph()
            # 한 항목이 두 페이지에 걸치지 않도록
            p.paragraph_format.keep_together = True
            if level == "h1":
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.4
                # Part 헤더는 다음 항목과 떨어지지 않게
                p.paragraph_format.keep_with_next = True
                _add_run(p, text, bold=True, color=BRAND_OR, size=11.5)
                _set_paragraph_border(p, bottom={"size": 6, "color": BRAND_OR, "space": 2})
            else:
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.6
                _add_run(p, "  " + text, color=TX_DARK, size=10)
            # 점선 + 번호 자리 (실제 페이지는 Word 가 위쪽 TOC 필드에서 채움)
            _add_run(p, "  " + ("·" * 8), color="BBBBBB", size=10)

    for el in soup.children:
        if not getattr(el, "name", None):
            continue
        name = el.name

        # 자동 목차 div → Word TOC 필드 + 정적 폴백
        if name == "div" and "auto-toc" in (el.get("class") or []):
            add_word_toc_field()
            add_static_toc_fallback()
            continue

        if name == "h1":
            if not first_h1:
                # 새 챕터 → 페이지 나눔
                br_p = doc.add_paragraph()
                br_p.add_run().add_break(WD_BREAK.PAGE)
            first_h1 = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            _add_run(p, el.get_text(), bold=True, color=BRAND_OR, size=20)
            _set_paragraph_border(p, bottom={"size": 16, "color": BRAND_OR})

        elif name == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.3)
            _add_run(p, el.get_text(), bold=True, color=TX_DARK, size=14)
            _set_paragraph_border(p, left={"size": 32, "color": BRAND_OR, "space": 8})

        elif name == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            _add_run(p, el.get_text(), bold=True, color=BRAND_OR_DARK, size=12)

        elif name == "h4":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            _add_run(p, el.get_text(), bold=True, color=TX_GRAY, size=11)

        elif name == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _walk_inline(p, el)

        elif name in ("ul", "ol"):
            style = "List Bullet" if name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(2)
                _walk_inline(p, li)

        elif name == "pre":
            code_text = el.get_text()
            # 어두운 배경 코드블록
            tbl = doc.add_table(rows=1, cols=1)
            tbl.autofit = True
            cell = tbl.cell(0, 0)
            _set_cell_shading(cell, CODE_BG_DARK)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            for i, line in enumerate(code_text.rstrip("\n").split("\n")):
                cp = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                cp.paragraph_format.line_spacing = 1.3
                cp.paragraph_format.space_after = Pt(0)
                _add_run(cp, line, color=CODE_FG_LIGHT, size=9, font=MONO_FONT)
            # 빈 단락으로 간격 확보
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)

        elif name == "blockquote":
            text = el.get_text().strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_run(p, text, color=TX_GRAY, size=10, italic=True)
            _set_paragraph_border(p, left={"size": 24, "color": BRAND_OR, "space": 8})
            _set_paragraph_shading(p, BRAND_OR_LIGHT)

        elif name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            ncols = max(len(r.find_all(["td", "th"])) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                is_header = any(c.name == "th" for c in cells)
                for ci, c in enumerate(cells):
                    if ci >= ncols:
                        break
                    cell = tbl.cell(ri, ci)
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    _add_run(
                        cp, c.get_text().strip(),
                        bold=is_header,
                        color=BRAND_OR_DARK if is_header else TX_DARK,
                        size=10,
                    )
                    if is_header:
                        _set_cell_shading(cell, BRAND_OR_LIGHT)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif name == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            _set_paragraph_border(p, bottom={"size": 12, "color": BRAND_OR})

        else:
            # 알 수 없는 블록 — 일반 단락으로
            text = el.get_text().strip()
            if text:
                p = doc.add_paragraph()
                _add_run(p, text)

    doc.save(str(DOCX))
    print(f"OK → {DOCX} ({DOCX.stat().st_size // 1024} KB) [styled]")


build_docx()
